import subprocess
import os
import sys
import datetime
import re
import csv
import platform
import gc

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.table import WD_ALIGN_VERTICAL
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

IS_WINDOWS = platform.system() == 'Windows'

def check_dataset_size(dataset_path):
    """检查数据集大小，预测是否可能失败"""
    try:
        file_size = os.path.getsize(dataset_path)
        num_lines = 0
        with open(dataset_path, 'r') as f:
            for line in f:
                if line.strip() and not line.startswith('#'):
                    num_lines += 1
        
        size_mb = file_size / (1024 * 1024)
        is_large = size_mb > 5 or num_lines > 1000000
        
        print(f"  📊 数据集大小: {size_mb:.2f} MB, {num_lines} 条边")
        
        if is_large:
            print(f"  ⚠️  数据集较大，可能内存不足")
        
        return is_large
    except Exception as e:
        print(f"  ⚠️  检查数据集大小失败: {e}")
        return False

def save_interim_results(all_results, step, csv_only=False):
    """保存临时结果，防止程序崩溃数据丢失"""
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    csv_file = f"临时结果_{step}_{timestamp}.csv"
    try:
        save_to_csv(all_results, csv_file)
        print(f"\n💾 已保存临时结果到 {csv_file}")
    except Exception as e:
        print(f"⚠️ 保存临时结果失败: {e}")
    return csv_file

def run_wsl_command(cmd, cwd, timeout=300):
    if IS_WINDOWS:
        full_cmd = ['wsl', '-e', 'bash', '-c', f'cd "{cwd}" && {" ".join(cmd)}']
    else:
        full_cmd = cmd
    
    print(f"命令: {' '.join(full_cmd)}")
    try:
        output = subprocess.run(full_cmd, capture_output=True, text=True, timeout=timeout)
        return output.stdout, output.stderr, output.returncode
    except subprocess.TimeoutExpired:
        print(f"⚠️ 命令执行超时 ({timeout}秒)")
        return "", "超时", -1
    except Exception as e:
        print(f"⚠️ 命令执行异常: {e}")
        return "", str(e), -1

def check_cugraph_installed():
    conda_python = '/home/dbh/miniconda3/envs/cugraph_env/bin/python'
    stdout, stderr, returncode = run_wsl_command(['ls', '-la', conda_python], '/mnt/d/Desktop/result/cugraph-main')
    if returncode != 0:
        print(f"  ⚠️ 环境不存在: {conda_python}")
        return False
    
    print("  ✅ 环境存在，跳过详细测试（可能会卡住）")
    return True

def run_cugraph_algorithm(dataset_path):
    result = {
        'algorithm': 'cuGraph SCC',
        'file_name': os.path.basename(dataset_path),
        'graph_type': '有向图',
        'num_vertices': 'N/A',
        'num_edges': 'N/A',
        'total_scc': 'N/A',
        'biggest_scc': 'N/A',
        'total_time': 'N/A',
        'read_time': 'N/A',
        'create_time': 'N/A',
        'compute_time': 'N/A',
        'rss_peak': 'N/A',
        'memory_usage': 'N/A',
        'algorithm_result': 'N/A',
        'success': False,
        'error_type': ''
    }
    
    try:
        file_size = os.path.getsize(dataset_path)
        size_mb = file_size / (1024 * 1024)
        
        # 根据数据集大小设置超时时间
        if size_mb > 50:
            timeout = 1800  # 30分钟
            print(f"  ⚠️ 大型数据集 ({size_mb:.1f} MB)，超时设置为 {timeout//60} 分钟")
        elif size_mb > 10:
            timeout = 1200  # 20分钟
            print(f"  ⚠️ 较大数据集 ({size_mb:.1f} MB)，超时设置为 {timeout//60} 分钟")
        elif size_mb > 5:
            timeout = 600  # 10分钟
            print(f"  ⚠️ 中等数据集 ({size_mb:.1f} MB)，超时设置为 {timeout//60} 分钟")
        else:
            timeout = 300  # 5分钟
            print(f"  数据集大小: {size_mb:.1f} MB，超时设置为 {timeout//60} 分钟")
        
        cwd = '/mnt/d/Desktop/result/cugraph-main'
        dataset_wsl_path = '/mnt/d/Desktop/' + os.path.relpath(dataset_path, 'd:/Desktop').replace('\\', '/')
        conda_python = '/home/dbh/miniconda3/envs/cugraph_env/bin/python'
        cmd = [conda_python, 'run_edge_list.py', dataset_wsl_path]
        stdout, stderr, returncode = run_wsl_command(cmd, cwd, timeout=timeout)
        
        output_text = stdout + stderr
        
        if returncode == 0:
            result['success'] = True
            print("✅ cuGraph SCC 运行成功")
            print(output_text)
            
            vertices_match = re.search(r'重映射后顶点数量[:=]\s*([\d.]+)', output_text)
            if vertices_match:
                result['num_vertices'] = vertices_match.group(1)
            
            edges_match = re.search(r'边列表形状[:=]\s*\(([\d.]+)', output_text)
            if edges_match:
                result['num_edges'] = edges_match.group(1)
            
            read_time_match = re.search(r'成功读取文件，耗时[:=]\s*([\d.]+)', output_text)
            if read_time_match:
                result['read_time'] = float(read_time_match.group(1))
            
            create_time_match = re.search(r'成功创建图，耗时[:=]\s*([\d.]+)', output_text)
            if create_time_match:
                result['create_time'] = float(create_time_match.group(1))
            
            scc_time_match = re.search(r'成功运行.*耗时[:=]\s*([\d.]+)', output_text)
            if scc_time_match:
                result['compute_time'] = float(scc_time_match.group(1))
            
            total_time = 0
            if result['read_time'] != 'N/A':
                total_time += result['read_time']
            if result['create_time'] != 'N/A':
                total_time += result['create_time']
            if result['compute_time'] != 'N/A':
                total_time += result['compute_time']
            
            if total_time > 0:
                result['total_time'] = f"{total_time:.4f}"
            
            gpu_mem_match = re.search(r'GPU显存[:=]\s*(\d+)', output_text)
            if gpu_mem_match:
                result['rss_peak'] = gpu_mem_match.group(1)
            
            scc_count_match = re.search(r'SCC总数[:=]\s*(\d+)', output_text)
            if scc_count_match:
                result['total_scc'] = scc_count_match.group(1)
            else:
                scc_count_match = re.search(r'总 SCC 数量[:=]\s*(\d+)', output_text)
                if scc_count_match:
                    result['total_scc'] = scc_count_match.group(1)
            
            biggest_match = re.search(r'最大SCC大小[:=]\s*(\d+)', output_text)
            if biggest_match:
                result['biggest_scc'] = biggest_match.group(1)
            else:
                biggest_match = re.search(r'最大 SCC 大小[:=]\s*(\d+)', output_text)
                if biggest_match:
                    result['biggest_scc'] = biggest_match.group(1)
            
            memory_match = re.search(r'内存使用[:=]\s*(\d+)', output_text)
            if memory_match:
                result['memory_usage'] = f"{memory_match.group(1)} MB"
            
            rss_peak_match = re.search(r'RSS峰值[:=]\s*(\d+)', output_text)
            if rss_peak_match:
                result['rss_peak'] = rss_peak_match.group(1)
            
            trivial_match = re.search(r'平凡SCC数[:=]\s*(\d+)', output_text)
            if trivial_match:
                result['num_trivial'] = trivial_match.group(1)
            
            nontrivial_match = re.search(r'非平凡SCC数[:=]\s*(\d+)', output_text)
            if nontrivial_match:
                result['num_nontrivial'] = nontrivial_match.group(1)
            
            algorithm_result_match = re.search(r'算法结果[:=]\s*(\S+)', output_text)
            if algorithm_result_match:
                result['algorithm_result'] = algorithm_result_match.group(1)
                
        else:
            result['error'] = stderr[:200] if stderr else stdout[:200]
            if 'ModuleNotFoundError' in result['error'] or 'No module named' in result['error']:
                result['error_type'] = 'missing_module'
            elif 'MemoryError' in result['error']:
                result['error_type'] = 'memory'
            else:
                result['error_type'] = 'other'
            print(f"❌ cuGraph 运行失败: {result['error']}")
            
    except Exception as e:
        result['error'] = str(e)
        result['error_type'] = 'exception'
        print(f"❌ cuGraph 运行异常: {e}")
    
    return result

def run_gardenia_algorithm(dataset_path):
    result = {
        'algorithm': 'Gardenia SCC (GPU)',
        'file_name': os.path.basename(dataset_path),
        'graph_type': '有向图',
        'num_vertices': 'N/A',
        'num_edges': 'N/A',
        'total_scc': 'N/A',
        'biggest_scc': 'N/A',
        'num_trivial': 'N/A',
        'num_nontrivial': 'N/A',
        'total_time': 'N/A',
        'read_time': 'N/A',
        'compute_time': 'N/A',
        'verify_time': 'N/A',
        'rss_peak': 'N/A',
        'memory_usage': 'N/A',
        'algorithm_result': 'N/A',
        'success': False,
        'error': ''
    }
    
    try:
        file_size = os.path.getsize(dataset_path)
        size_mb = file_size / (1024 * 1024)
        
        # 根据数据集大小设置超时时间
        if size_mb > 50:
            timeout = 1800  # 30分钟
        elif size_mb > 10:
            timeout = 1200  # 20分钟
        elif size_mb > 5:
            timeout = 600  # 10分钟
        else:
            timeout = 300  # 5分钟
        
        dataset_wsl_path = '/mnt/d/Desktop/' + os.path.relpath(dataset_path, 'd:/Desktop').replace('\\', '/')
        conda_python = '/home/dbh/miniconda3/envs/cugraph_env/bin/python'
        gardenia_gpu_script = '/mnt/d/Desktop/result/gardenia_scc_gpu.py'
        cmd = [conda_python, gardenia_gpu_script, dataset_wsl_path]
        cwd = '/mnt/d/Desktop/result'
        stdout, stderr, returncode = run_wsl_command(cmd, cwd, timeout=timeout)
        
        if returncode == 0:
            result['success'] = True
            print("✅ Gardenia 运行成功")
            
            output_text = stdout + stderr
            print(output_text)
            
            vertices_match = re.search(r'顶点数[:=]\s*(\d+)', output_text)
            if vertices_match:
                result['num_vertices'] = vertices_match.group(1)
            
            edges_match = re.search(r'边数[:=]\s*(\d+)', output_text)
            if edges_match:
                result['num_edges'] = edges_match.group(1)
            
            total_scc_match = re.search(r'SCC总数[:=]\s*(\d+)', output_text)
            if total_scc_match:
                result['total_scc'] = total_scc_match.group(1)
            
            biggest_match = re.search(r'最大SCC大小[:=]\s*(\d+)', output_text)
            if biggest_match:
                result['biggest_scc'] = biggest_match.group(1)
            
            trivial_match = re.search(r'平凡SCC数[:=]\s*(\d+)', output_text)
            if trivial_match:
                result['num_trivial'] = trivial_match.group(1)
            
            nontrivial_match = re.search(r'非平凡SCC数[:=]\s*(\d+)', output_text)
            if nontrivial_match:
                result['num_nontrivial'] = nontrivial_match.group(1)
            
            compute_time_match = re.search(r'计算时间[:=]\s*([\d.]+)', output_text)
            if compute_time_match:
                result['compute_time'] = float(compute_time_match.group(1))
            
            rss_peak_match = re.search(r'RSS峰值[:=]\s*(\d+)', output_text)
            if rss_peak_match:
                result['rss_peak'] = rss_peak_match.group(1)
                result['memory_usage'] = f"{rss_peak_match.group(1)} MB"
            
            total_time_match = re.search(r'总时间[:=]\s*([\d.]+)', output_text)
            if total_time_match:
                result['total_time'] = f"{float(total_time_match.group(1)):.4f}"
            
            result['algorithm_result'] = '成功'
            
            if result['read_time'] != 'N/A' or result['compute_time'] != 'N/A':
                total_ms = 0
                if result['read_time'] != 'N/A':
                    total_ms += result['read_time']
                if result['compute_time'] != 'N/A':
                    total_ms += result['compute_time']
                if result['verify_time'] != 'N/A':
                    total_ms += result['verify_time']
                if total_ms > 0:
                    result['total_time'] = f"{total_ms:.4f}"
                
        else:
            result['error'] = stderr[:200]
            print(f"❌ Gardenia 运行失败: {result['error']}")
            
    except Exception as e:
        result['error'] = str(e)
        print(f"❌ Gardenia 运行异常: {e}")
    
    return result

def generate_table(results):
    if len(results) == 0:
        return "无测试结果"
    
    print("\n" + "="*140)
    print("                    SCC 算法对比结果表格（相同数据集）")
    print("="*140)
    
    headers = ["数据集", "算法", "顶点数", "边数", "SCC总数", "最大SCC", "总时间", "读取时间", "计算时间", "内存使用", "状态"]
    
    col_widths = [18, 15, 10, 10, 10, 12, 12, 12, 12, 14, 10]
    
    header_line = "|"
    for i, header in enumerate(headers):
        header_line += f" {header:{col_widths[i]}} |"
    print(header_line)
    print("|" + "-"*(sum(col_widths) + len(col_widths)*3 - 1) + "|")
    
    for result in results:
        status = "✓ 成功" if result['success'] else f"✗ 失败"
        read_time = f"{result.get('read_time', 'N/A'):.4f}" if result.get('read_time') != 'N/A' else "N/A"
        compute_time = f"{result.get('compute_time', 'N/A'):.4f}" if result.get('compute_time') != 'N/A' else "N/A"
        
        memory_usage = result.get('memory_usage', 'N/A')
        if memory_usage == 'N/A':
            memory_usage = f"{result.get('rss_peak', 'N/A')} MB" if result.get('rss_peak', 'N/A') != 'N/A' else "N/A"
        
        rss_peak = result.get('rss_peak', 'N/A')
        if rss_peak != 'N/A':
            rss_peak_display = f"{rss_peak} MB"
        else:
            rss_peak_display = "N/A"
        
        row = f"| {result['file_name']:{col_widths[0]}} |"
        row += f" {result['algorithm']:{col_widths[1]}} |"
        row += f" {result['num_vertices']:{col_widths[2]}} |"
        row += f" {result['num_edges']:{col_widths[3]}} |"
        row += f" {result['total_scc']:{col_widths[4]}} |"
        row += f" {result['biggest_scc']:{col_widths[5]}} |"
        row += f" {result['total_time']:{col_widths[6]}} |"
        row += f" {read_time:{col_widths[7]}} |"
        row += f" {compute_time:{col_widths[8]}} |"
        row += f" {memory_usage:{col_widths[9]}} |"
        row += f" {status:{col_widths[10]}} |"
        print(row)
    
    print("="*140)
    
    return headers, results

def save_to_csv(results, output_file):
    headers = ["数据集", "算法", "图类型", "顶点数", "边数", "SCC总数", "最大SCC", "平凡SCC", "非平凡SCC", "总时间", "读取时间", "计算时间", "验证时间", "内存使用", "内存使用峰值", "算法结果", "状态", "错误信息"]
    
    with open(output_file, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow(headers)
        
        for result in results:
            status = "成功" if result['success'] else "失败"
            
            read_time_val = result.get('read_time', 'N/A')
            read_time = f"{read_time_val:.4f}" if isinstance(read_time_val, (int, float)) else str(read_time_val)
            
            compute_time_val = result.get('compute_time', 'N/A')
            compute_time = f"{compute_time_val:.4f}" if isinstance(compute_time_val, (int, float)) else str(compute_time_val)
            
            verify_time_val = result.get('verify_time', 'N/A')
            verify_time = f"{verify_time_val:.4f}" if isinstance(verify_time_val, (int, float)) else str(verify_time_val)
            
            memory_usage = result.get('memory_usage', 'N/A')
            if memory_usage == 'N/A':
                memory_usage = f"{result.get('rss_peak', 'N/A')} MB" if result.get('rss_peak', 'N/A') != 'N/A' else "N/A"
            
            rss_peak = result.get('rss_peak', 'N/A')
            if rss_peak != 'N/A':
                rss_peak_display = f"{rss_peak} MB"
            else:
                rss_peak_display = "N/A"
            
            row = [
                result['file_name'],
                result['algorithm'],
                result.get('graph_type', '有向图'),
                result['num_vertices'],
                result['num_edges'],
                result['total_scc'],
                result['biggest_scc'],
                result.get('num_trivial', 'N/A'),
                result.get('num_nontrivial', 'N/A'),
                result['total_time'],
                read_time,
                compute_time,
                verify_time,
                memory_usage,
                rss_peak_display,
                result.get('algorithm_result', 'N/A'),
                status,
                result.get('error', '')
            ]
            writer.writerow(row)
    
    print(f"\n📊 结果已保存到 {output_file}")

def save_to_word(results, output_file, cugraph_available):
    if not DOCX_AVAILABLE:
        print("⚠️ 无法保存Word文件：需要安装 python-docx 库")
        print("请运行: pip install python-docx")
        return False
    
    try:
        doc = Document()
        
        title = doc.add_heading('SCC 算法对比测试报告', 0)
        title.alignment = 1
        
        doc.add_paragraph(f"生成日期: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"测试算法数量: 2 (cuGraph, Gardenia)")
        doc.add_paragraph(f"测试数据集数量: {len(results) // 2}")
        doc.add_paragraph(f"运行环境: {'Windows + WSL' if IS_WINDOWS else 'Linux/WSL'}")
        doc.add_paragraph(f"cuGraph状态: {'已安装' if cugraph_available else '未安装'}")
        
        doc.add_heading('测试结果对比表（相同数据集）', level=1)
        
        table = doc.add_table(rows=1, cols=12)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        headers = ["数据集", "算法", "顶点数", "边数", "SCC总数", "最大SCC", "总时间", "读取时间", "计算时间", "内存使用", "内存峰值", "状态"]
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        for result in results:
            row_cells = table.add_row().cells
            read_time = f"{result.get('read_time', 'N/A'):.4f}" if result.get('read_time') != 'N/A' else "N/A"
            compute_time = f"{result.get('compute_time', 'N/A'):.4f}" if result.get('compute_time') != 'N/A' else "N/A"
            
            row_cells[0].text = result['file_name']
            row_cells[1].text = result['algorithm']
            row_cells[2].text = result['num_vertices']
            row_cells[3].text = result['num_edges']
            row_cells[4].text = result['total_scc']
            row_cells[5].text = result['biggest_scc']
            row_cells[6].text = result['total_time']
            row_cells[7].text = read_time
            row_cells[8].text = compute_time
            row_cells[9].text = result.get('memory_usage', 'N/A')
            row_cells[10].text = result.get('rss_peak', 'N/A')
            row_cells[11].text = "成功" if result['success'] else "失败"
            
            for cell in row_cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        doc.add_heading('算法性能对比分析', level=1)
        
        datasets = {}
        for result in results:
            key = result['file_name']
            if key not in datasets:
                datasets[key] = {'gardenia': None, 'cugraph': None}
            if result['algorithm'] == 'Gardenia SCC (GPU)':
                datasets[key]['gardenia'] = result
            else:
                datasets[key]['cugraph'] = result
        
        for dataset_name, data in datasets.items():
            doc.add_heading(f"数据集: {dataset_name}", level=2)
            
            gardenia = data['gardenia']
            cugraph = data['cugraph']
            
            if gardenia and gardenia['success']:
                doc.add_paragraph(f"**Gardenia SCC**:")
                doc.add_paragraph(f"  - 图类型: {gardenia.get('graph_type', '有向图')}")
                doc.add_paragraph(f"  - 顶点数: {gardenia['num_vertices']}")
                doc.add_paragraph(f"  - 边数: {gardenia['num_edges']}")
                doc.add_paragraph(f"  - SCC总数: {gardenia['total_scc']}")
                doc.add_paragraph(f"  - 最大SCC: {gardenia['biggest_scc']}")
                doc.add_paragraph(f"  - 平凡SCC: {gardenia['num_trivial']}")
                doc.add_paragraph(f"  - 非平凡SCC: {gardenia['num_nontrivial']}")
                
                read_time = gardenia.get('read_time', 'N/A')
                read_time_str = f"{read_time:.4f}" if isinstance(read_time, (int, float)) else str(read_time)
                doc.add_paragraph(f"  - 读取时间: {read_time_str} 秒")
                
                compute_time = gardenia.get('compute_time', 'N/A')
                compute_time_str = f"{compute_time:.4f}" if isinstance(compute_time, (int, float)) else str(compute_time)
                doc.add_paragraph(f"  - 计算时间: {compute_time_str} 秒")
                
                verify_time = gardenia.get('verify_time', 'N/A')
                verify_time_str = f"{verify_time:.4f}" if isinstance(verify_time, (int, float)) else str(verify_time)
                doc.add_paragraph(f"  - 验证时间: {verify_time_str} 秒")
                
                doc.add_paragraph(f"  - 总时间: {gardenia['total_time']} 秒")
                doc.add_paragraph(f"  - 内存使用峰值: {gardenia.get('memory_usage', 'N/A')}")
                doc.add_paragraph(f"  - 算法结果: {gardenia.get('algorithm_result', 'N/A')}")
            else:
                doc.add_paragraph(f"**Gardenia SCC**: ❌ 未成功运行")
            
            if cugraph and cugraph['success']:
                doc.add_paragraph(f"**cuGraph SCC**:")
                doc.add_paragraph(f"  - 图类型: {cugraph.get('graph_type', '有向图')}")
                doc.add_paragraph(f"  - 顶点数: {cugraph['num_vertices']}")
                doc.add_paragraph(f"  - 边数: {cugraph['num_edges']}")
                doc.add_paragraph(f"  - SCC总数: {cugraph['total_scc']}")
                doc.add_paragraph(f"  - 最大SCC: {cugraph['biggest_scc']}")
                
                cugraph_read_time = cugraph.get('read_time', 'N/A')
                cugraph_read_time_str = f"{cugraph_read_time:.4f}" if isinstance(cugraph_read_time, (int, float)) else str(cugraph_read_time)
                doc.add_paragraph(f"  - 读取时间: {cugraph_read_time_str} 秒")
                
                cugraph_create_time = cugraph.get('create_time', 'N/A')
                cugraph_create_time_str = f"{cugraph_create_time:.4f}" if isinstance(cugraph_create_time, (int, float)) else str(cugraph_create_time)
                doc.add_paragraph(f"  - 建图时间: {cugraph_create_time_str} 秒")
                
                cugraph_compute_time = cugraph.get('compute_time', 'N/A')
                cugraph_compute_time_str = f"{cugraph_compute_time:.4f}" if isinstance(cugraph_compute_time, (int, float)) else str(cugraph_compute_time)
                doc.add_paragraph(f"  - 计算时间: {cugraph_compute_time_str} 秒")
                
                doc.add_paragraph(f"  - 总时间: {cugraph['total_time']} 秒")
                doc.add_paragraph(f"  - 内存使用: {cugraph.get('memory_usage', 'N/A')}")
                doc.add_paragraph(f"  - 内存使用峰值: {cugraph.get('rss_peak', 'N/A')} MB")
                doc.add_paragraph(f"  - 算法结果: {cugraph.get('algorithm_result', 'N/A')}")
                
                if gardenia and gardenia['success'] and gardenia['total_time'] != 'N/A' and cugraph['total_time'] != 'N/A':
                    gardenia_time = float(gardenia['total_time'])
                    cugraph_time = float(cugraph['total_time'])
                    if cugraph_time > 0:
                        speedup = gardenia_time / cugraph_time
                        doc.add_paragraph(f"**性能对比**:")
                        doc.add_paragraph(f"  - Gardenia SCC 耗时: {gardenia_time:.4f} 秒")
                        doc.add_paragraph(f"  - cuGraph SCC 耗时: {cugraph_time:.4f} 秒")
                        doc.add_paragraph(f"  - {'Gardenia SCC' if speedup < 1 else 'cuGraph SCC'} 快 {max(speedup, 1/speedup):.2f} 倍")
            else:
                doc.add_paragraph(f"**cuGraph SCC**: ❌ 未成功运行")
        
        doc.add_heading('测试环境信息', level=1)
        doc.add_paragraph("- 操作系统: Windows 10/11 + WSL2")
        doc.add_paragraph("- 算法1: cuGraph SCC (GPU加速)")
        doc.add_paragraph("- 算法2: Gardenia SCC (GPU，cuGraph实现)")
        doc.add_paragraph("- 数据集格式: 边列表格式 (.txt)")
        
        doc.add_heading('注意事项', level=1)
        doc.add_paragraph("1. cuGraph需要NVIDIA GPU和CUDA支持")
        doc.add_paragraph("2. Gardenia是Linux程序，通过WSL在Windows上运行")
        doc.add_paragraph("3. 两个算法运行相同的数据集进行对比")
        doc.add_paragraph("4. 顶点数差异说明:")
        doc.add_paragraph("   - Gardenia: 统计所有出现的顶点ID的最大值+1（包含孤立顶点）")
        doc.add_paragraph("   - cuGraph: 只统计实际有边的顶点（进行了重映射）")
        doc.add_paragraph("   - 这是算法设计差异，非错误")
        
        doc.save(output_file)
        print(f"\n📄 Word报告已保存到 {output_file}")
        return True
    except Exception as e:
        print(f"❌ 保存Word文件失败: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    test_folder = 'd:/Desktop/test'
    
    default_datasets = []
    if os.path.isdir(test_folder):
        for file in os.listdir(test_folder):
            if file.endswith('.txt'):
                file_path = os.path.join(test_folder, file)
                try:
                    file_size = os.path.getsize(file_path)
                    default_datasets.append((file_size, file_path))
                except:
                    default_datasets.append((0, file_path))
        # 按文件大小排序，优先处理小数据集
        default_datasets.sort()
        default_datasets = [path for (size, path) in default_datasets]
    
    if not default_datasets:
        default_datasets = [
            'd:/Desktop/test/edge_list.txt',
            'd:/Desktop/test/web-Google.txt'
        ]
    
    datasets = []
    
    print("="*120)
    print("                    SCC 算法对比测试脚本（相同数据集对比）")
    print("="*120)
    print(f"测试日期: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"运行环境: {'Windows + WSL' if IS_WINDOWS else 'Linux/WSL'}")
    print("="*120)
    
    if len(sys.argv) > 1:
        print("\n📋 使用命令行指定的数据集...")
        args = sys.argv[1:]
        
        for arg in args:
            if os.path.exists(arg):
                datasets.append(arg)
                print(f"  ✓ 添加数据集: {os.path.basename(arg)}")
            else:
                print(f"  ⚠️ 文件不存在: {arg}，跳过")
        
        if len(datasets) == 0:
            print("  ⚠️ 没有有效的数据集，使用默认数据集")
            datasets = default_datasets
    else:
        print("\n📋 使用默认数据集 (桌面test文件夹)")
        datasets = default_datasets
    
    print(f"\n测试数据集数量: {len(datasets)}")
    print("="*120)
    
    print("\n📋 检查环境...")
    cugraph_available = check_cugraph_installed()
    print(f"cuGraph状态: {'✅ 已安装' if cugraph_available else '❌ 未安装'}")
    
    all_results = []
    successful_results = 0
    failed_results = 0
    
    for i, dataset_path in enumerate(datasets):
        dataset_name = os.path.basename(dataset_path)
        
        # 检查数据集大小
        try:
            file_size = os.path.getsize(dataset_path)
            size_mb = file_size / (1024 * 1024)
            num_lines = 0
            try:
                with open(dataset_path, 'r') as f:
                    for line in f:
                        if line.strip() and not line.startswith('#'):
                            num_lines += 1
            except:
                num_lines = 0
        except:
            size_mb = 0
            num_lines = 0
        
        print(f"\n{'='*80}")
        print(f"正在测试数据集 #{i+1}/{len(datasets)}: {dataset_name}")
        print(f"  📊 大小: {size_mb:.2f} MB, 边数: {num_lines}")
        print(f"{'='*80}")
        
        # 对于超大数据集，标记可能的风险
        skip_large = False
        if size_mb > 50:
            print(f"⚠️ 数据集非常大，可能会失败，但仍会尝试...")
        elif size_mb > 10:
            print(f"⚠️ 数据集较大，可能需要更多内存")
        
        dataset_success = 0
        dataset_failed = 0
        
        try:
            print("\n--- Gardenia SCC ---")
            gardenia_result = None
            try:
                gardenia_result = run_gardenia_algorithm(dataset_path)
                all_results.append(gardenia_result)
                
                if gardenia_result['success']:
                    successful_results += 1
                    dataset_success += 1
                    print(f"✅ Gardenia SCC 成功完成！")
                else:
                    failed_results += 1
                    dataset_failed += 1
                    print(f"❌ Gardenia SCC 失败: {gardenia_result.get('error', '未知错误')}")
            except Exception as ge:
                print(f"❌ Gardenia 异常: {ge}")
                # 添加失败的结果记录
                failed_gardenia = {
                    'algorithm': 'Gardenia SCC (GPU)',
                    'file_name': dataset_name,
                    'graph_type': '有向图',
                    'num_vertices': 'N/A',
                    'num_edges': 'N/A',
                    'total_scc': 'N/A',
                    'biggest_scc': 'N/A',
                    'num_trivial': 'N/A',
                    'num_nontrivial': 'N/A',
                    'total_time': 'N/A',
                    'read_time': 'N/A',
                    'compute_time': 'N/A',
                    'verify_time': 'N/A',
                    'rss_peak': 'N/A',
                    'memory_usage': 'N/A',
                    'algorithm_result': '失败',
                    'success': False,
                    'error': str(ge)
                }
                all_results.append(failed_gardenia)
                failed_results += 1
                dataset_failed += 1
            
            gc.collect()
            import time
            time.sleep(2)  # 给系统清理内存的时间
            print("  🧹 Gardenia内存已清理")
            
            print("\n--- cuGraph SCC ---")
            cugraph_result = None
            try:
                cugraph_result = run_cugraph_algorithm(dataset_path)
                all_results.append(cugraph_result)
                
                if cugraph_result['success']:
                    successful_results += 1
                    dataset_success += 1
                    print(f"✅ cuGraph SCC 成功完成！")
                else:
                    failed_results += 1
                    dataset_failed += 1
                    print(f"❌ cuGraph 失败: {cugraph_result.get('error', '未知错误')}")
            except Exception as ce:
                print(f"❌ cuGraph 异常: {ce}")
                # 添加失败的结果记录
                failed_cugraph = {
                    'algorithm': 'cuGraph SCC',
                    'file_name': dataset_name,
                    'graph_type': '有向图',
                    'num_vertices': 'N/A',
                    'num_edges': 'N/A',
                    'total_scc': 'N/A',
                    'biggest_scc': 'N/A',
                    'num_trivial': 'N/A',
                    'num_nontrivial': 'N/A',
                    'total_time': 'N/A',
                    'read_time': 'N/A',
                    'create_time': 'N/A',
                    'compute_time': 'N/A',
                    'rss_peak': 'N/A',
                    'memory_usage': 'N/A',
                    'algorithm_result': '失败',
                    'success': False,
                    'error': str(ce)
                }
                all_results.append(failed_cugraph)
                failed_results += 1
                dataset_failed += 1
            
            gc.collect()
            time.sleep(2)  # 给系统清理内存的时间
            print("  🧹 cuGraph内存已清理")
        
        except Exception as e:
            print(f"\n❌ 处理 {dataset_name} 时发生错误: {e}")
            import traceback
            traceback.print_exc()
            # 记录失败结果，继续下一个
            failed_result = {
                'algorithm': '处理失败',
                'file_name': dataset_name,
                'graph_type': '有向图',
                'num_vertices': 'N/A',
                'num_edges': 'N/A',
                'total_scc': 'N/A',
                'biggest_scc': 'N/A',
                'num_trivial': 'N/A',
                'num_nontrivial': 'N/A',
                'total_time': 'N/A',
                'read_time': 'N/A',
                'compute_time': 'N/A',
                'verify_time': 'N/A',
                'rss_peak': 'N/A',
                'memory_usage': 'N/A',
                'algorithm_result': '失败',
                'success': False,
                'error': str(e)
            }
            all_results.append(failed_result)
            failed_results += 2
        
        # 每处理完一个数据集就保存一次
        print(f"\n📊 数据集 {dataset_name} 完成: 成功 {dataset_success}/2, 失败 {dataset_failed}/2")
        print(f"📊 整体进度: {i+1}/{len(datasets)} 数据集, 成功: {successful_results}, 失败: {failed_results}")
        
        if (i + 1) % 1 == 0:  # 每完成一个就保存
            save_interim_results(all_results, f"{i+1}of{len(datasets)}")
            gc.collect()
    
    # 生成最终报告
    print(f"\n{'='*80}")
    print("                    最终测试统计")
    print(f"{'='*80}")
    print(f"总计处理: {len(datasets)} 个数据集")
    print(f"成功结果: {successful_results}")
    print(f"失败结果: {failed_results}")
    print(f"成功率: {(successful_results / max(len(all_results), 1) * 100):.1f}%")
    print(f"{'='*80}")
    
    generate_table(all_results)
    
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    
    csv_file = f"算法对比结果_{timestamp}.csv"
    save_to_csv(all_results, csv_file)
    
    if DOCX_AVAILABLE:
        word_file = f"算法对比报告_{timestamp}.docx"
        save_to_word(all_results, word_file, cugraph_available)
    else:
        print("\n⚠️ Word报告生成跳过: python-docx未安装")
        print(f"   CSV结果已保存到: {csv_file}")
    
    print("\n🎉 测试完成！")
    print("\n📖 使用说明:")
    print("  python run_both_algorithms.py                          # 使用默认数据集")
    print("  python run_both_algorithms.py dataset.txt              # 指定单个数据集")
    print("  python run_both_algorithms.py data1.txt data2.txt     # 指定多个数据集")
    if not cugraph_available:
        print("\n💡 提示: 要运行cuGraph，请安装NVIDIA RAPIDS环境")
